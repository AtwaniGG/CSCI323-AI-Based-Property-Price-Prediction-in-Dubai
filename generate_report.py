from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

# --- Style setup ---
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.5

for level in range(1, 4):
    h_style = doc.styles[f'Heading {level}']
    h_style.font.name = 'Times New Roman'
    h_style.font.color.rgb = RGBColor(0, 0, 0)
    if level == 1:
        h_style.font.size = Pt(16)
        h_style.font.bold = True
    elif level == 2:
        h_style.font.size = Pt(14)
        h_style.font.bold = True
    else:
        h_style.font.size = Pt(12)
        h_style.font.bold = True

def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)
                r.font.name = 'Times New Roman'
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
                    r.font.name = 'Times New Roman'
    doc.add_paragraph()
    return table

def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        p.add_run(text).font.name = 'Times New Roman'
        p.runs[-1].font.size = Pt(12)
    else:
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)

def add_numbered(doc, text):
    p = doc.add_paragraph(style='List Number')
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

# ============================================================
# TITLE PAGE
# ============================================================
for _ in range(6):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('AI-Driven Real Estate Price Prediction\nfor the Dubai Property Market')
run.bold = True
run.font.size = Pt(24)
run.font.name = 'Times New Roman'

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('CSCI323 – Modern Artificial Intelligence')
run.font.size = Pt(14)
run.font.name = 'Times New Roman'

doc.add_paragraph()

info_lines = [
    'Team Members:',
    '[Name 1] – [Student ID]',
    '[Name 2] – [Student ID]',
    '[Name 3] – [Student ID]',
    '[Name 4] – [Student ID]',
    '',
    'Submission Date: [Date]',
    'Instructor: [Instructor Name]',
]
for line in info_lines:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(line)
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'

doc.add_page_break()

# ============================================================
# EXECUTIVE SUMMARY
# ============================================================
doc.add_heading('Executive Summary', level=1)

doc.add_paragraph(
    'This project presents an artificial intelligence solution for predicting residential property '
    'prices per square meter (AED/sqm) in the Dubai real estate market. Using a dataset of over 1.6 million '
    'historical transactions from the Dubai Land Department, the system applies distributed big data processing '
    'via Apache Spark and deep learning via PyTorch to deliver accurate pricing forecasts.'
)

doc.add_paragraph(
    'Four modeling approaches were implemented and compared: a baseline Decision Tree, a Bagging Ensemble, '
    'a Gradient Boosting ensemble (all built from scratch using Spark ML), and a Multi-Layer Perceptron (MLP) '
    'Neural Network built with PyTorch. After extensive feature engineering producing 106 predictive features '
    'and rigorous evaluation, the Bagging Ensemble achieved the lowest RMSE among Spark-based models at '
    '3,273.04 AED/sqm, while the Neural Network MLP achieved an R² score of 0.7625 with an RMSE of '
    '3,810.88 AED/sqm and MAE of 2,735.13 AED/sqm.'
)

doc.add_paragraph(
    'These results demonstrate that AI-driven property valuation can provide reliable pricing guidance '
    'for investors, developers, and regulators in one of the world\'s most dynamic property markets.'
)

doc.add_page_break()

# ============================================================
# TABLE OF CONTENTS
# ============================================================
doc.add_heading('Table of Contents', level=1)

toc_items = [
    '1. Introduction and Motivation',
    '2. Literature Review',
    '3. Technical Requirements and Analysis',
    '4. Solution Methodology and Implementation',
    '5. Results and Evaluation',
    '6. Limitations, Challenges, and Constraints',
    '7. Conclusions and Discussion',
    '8. Recommendations and Future Work',
    '9. References',
    '10. Appendices',
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)

doc.add_paragraph()
doc.add_heading('List of Tables', level=2)

tables_list = [
    'Table 1: Dataset Summary Statistics',
    'Table 2: Feature Categories and Counts',
    'Table 3: Spark Model 10-Fold CV Results',
    'Table 4: Neural Network MLP Evaluation Metrics',
    'Table 5: Final Model Comparison',
    'Table 6: Objective Achievement Summary',
]
for t in tables_list:
    p = doc.add_paragraph(t)
    p.paragraph_format.space_after = Pt(2)

doc.add_page_break()

# ============================================================
# 1. INTRODUCTION AND MOTIVATION
# ============================================================
doc.add_heading('1. Introduction and Motivation', level=1)

doc.add_heading('1.1 Industry Background', level=2)
doc.add_paragraph(
    'Dubai\'s real estate market is one of the most active globally. The Dubai Land Department has recorded '
    'over 1.6 million property transactions spanning residential, commercial, and industrial segments. The '
    'market is characterized by high volatility, rapid development cycles, and significant price variation '
    'across neighborhoods and property types.'
)
doc.add_paragraph(
    'Property valuation in Dubai has traditionally relied on manual appraisals and comparative market analyses. '
    'These methods are labor-intensive, subjective, and often lag behind market conditions. The sheer volume '
    'of transactions makes manual analysis at scale infeasible.'
)

doc.add_heading('1.2 Problem Statement', level=2)
doc.add_paragraph(
    'Accurately predicting property prices per square meter in Dubai is a complex challenge due to the '
    'interplay of numerous factors: location, property type, bedrooms, proximity to landmarks and metro '
    'stations, transaction type, and temporal market trends. Simple statistical models fail to capture the '
    'non-linear relationships inherent in real estate pricing.'
)
doc.add_paragraph(
    'The central question: Can AI models, trained on historical transaction data, predict residential '
    'property prices per square meter with sufficient accuracy to support automated valuation and market analysis?'
)

doc.add_heading('1.3 Why AI Is the Appropriate Solution', level=2)
doc.add_paragraph('The problem is well-suited for AI due to:')
add_bullet(doc, 'Prices depend on dozens of interacting features, making manual modeling impractical.', 'High dimensionality: ')
add_bullet(doc, 'A penthouse in a prime area behaves differently from a studio in the same location.', 'Non-linear relationships: ')
add_bullet(doc, '1.6 million transactions supports training complex models with distributed processing.', 'Large-scale data: ')
add_bullet(doc, 'AI can discover latent patterns that human analysts may overlook.', 'Pattern recognition: ')

doc.add_heading('1.4 Business Objectives and Success Metrics', level=2)
add_numbered(doc, 'Build an end-to-end pipeline from raw data to trained models.')
add_numbered(doc, 'Compare traditional ensemble methods with deep learning for tabular real estate data.')
add_numbered(doc, 'Achieve an R² score above 0.70.')
add_numbered(doc, 'Minimize RMSE and MAE for practically useful predictions.')
add_numbered(doc, 'Ensure reproducibility through Docker containerization.')

doc.add_page_break()

# ============================================================
# 2. LITERATURE REVIEW
# ============================================================
doc.add_heading('2. Literature Review', level=1)

doc.add_paragraph(
    'Hedonic pricing models, introduced by Rosen (1974), established the foundation for property valuation '
    'by decomposing prices into constituent characteristics. However, these linear models struggle with '
    'non-linear interactions. Limsombunchai (2004) demonstrated that Artificial Neural Networks outperform '
    'hedonic models for housing price prediction when relationships are non-linear.'
)
doc.add_paragraph(
    'Gradient Boosting, introduced by Friedman (2001), builds ensembles of weak learners sequentially, '
    'with each tree correcting residual errors. This approach has proven highly effective for structured '
    'data (Chen & Guestrin, 2016). While libraries like XGBoost and LightGBM offer optimized implementations, '
    'building the algorithm from scratch provides deeper algorithmic understanding.'
)
doc.add_paragraph(
    'For neural networks on tabular data, Kadra et al. (2021) showed that well-regularized MLPs can match '
    'tree-based methods when appropriate techniques such as dropout and proper scaling are employed.'
)
doc.add_paragraph(
    'Apache Spark has become the standard for large-scale data processing and machine learning (Zaharia et al., '
    '2016). Spark MLlib provides distributed algorithm implementations, while PySpark enables Python-based '
    'workflows at scale.'
)
doc.add_paragraph(
    'Selim (2009) and Kok et al. (2014) explored models for Middle Eastern real estate markets, identifying '
    'location, size, and amenity proximity as dominant price determinants. Our work extends this by applying '
    'modern AI techniques to the comprehensive Dubai Land Department dataset at a scale not previously addressed.'
)

doc.add_page_break()

# ============================================================
# 3. TECHNICAL REQUIREMENTS
# ============================================================
doc.add_heading('3. Technical Requirements and Analysis', level=1)

doc.add_heading('3.1 System Requirements', level=2)
add_table(doc,
    ['Requirement', 'Specification'],
    [
        ['Processing Framework', 'Apache Spark 3.5.0'],
        ['Deep Learning Framework', 'PyTorch (GPU/CPU)'],
        ['Runtime Environment', 'Docker (linux/amd64)'],
        ['Memory', 'Minimum 4 GB RAM, 8 GB swap'],
        ['Python Version', 'Python 3.x'],
    ]
)

doc.add_heading('3.2 Data Requirements', level=2)
doc.add_paragraph(
    'The project uses the Dubai Land Department Transactions dataset containing 1,644,775 records across '
    '46 columns (279 MB CSV). Features include property type, sub-type, area, bedrooms, parking, location, '
    'proximity to landmarks/metro/malls, transaction type, and date. The target variable is actual_worth, '
    'transformed to price per square meter.'
)

doc.add_heading('3.3 Scalability Considerations', level=2)
add_bullet(doc, 'Docker containerization ensures consistent execution across platforms.')
add_bullet(doc, 'Apache Spark distributes computation and can scale to cluster deployments.')
add_bullet(doc, 'Modular pipeline: Preprocessing and modeling are decoupled for independent updates.')
add_bullet(doc, 'Checkpoint mechanism handles large joins without memory overflow.')

doc.add_page_break()

# ============================================================
# 4. METHODOLOGY AND IMPLEMENTATION
# ============================================================
doc.add_heading('4. Solution Methodology and Implementation', level=1)

doc.add_heading('4.1 AI Techniques Selected', level=2)
doc.add_paragraph('Four modeling approaches were selected:')
add_numbered(doc, 'Baseline Decision Tree (Spark ML): Reference point for ensemble improvements. Handles mixed data types natively.')
add_numbered(doc, 'Bagging Ensemble (from scratch, Spark): Trains multiple trees on random data subsets and averages predictions, reducing variance.')
add_numbered(doc, 'Gradient Boosting (from scratch, Spark): Trains trees sequentially on residual errors, reducing bias. Configuration: 10 estimators, learning rate 0.1, max depth 5.')
add_numbered(doc, 'Neural Network MLP (PyTorch): Three hidden layers to capture complex non-linear feature interactions.')

doc.add_heading('4.2 Data Preprocessing', level=2)
doc.add_paragraph('The raw dataset undergoes the following transformations using PySpark:')
add_numbered(doc, 'Bilingual cleanup: Arabic columns (suffix _ar) dropped; English suffixes (_en) removed.')
add_numbered(doc, 'ID column removal: Redundant IDs dropped in favor of human-readable names.')
add_numbered(doc, 'Procedure grouping: Raw procedure names simplified into Sales, Mortgage, Grants, Lease-to-Own, and Other.')
add_numbered(doc, 'Date parsing: instance_date parsed into sale_year and sale_quarter.')
add_numbered(doc, 'Bedroom extraction: Heterogeneous rooms values ("STUDIO", "3 B/R", "PENTHOUSE") parsed into numeric bedrooms and binary is_penthouse.')
add_numbered(doc, 'Property usage encoding: Decomposed into binary flags (residential, commercial, industrial, etc.).')
add_numbered(doc, 'Null handling: Defaults applied \u2014 master_project \u2192 "Standalone", nearest_landmark \u2192 "No Landmark", nearest_metro \u2192 "No Metro Access", nearest_mall \u2192 "No Major Mall".')
add_numbered(doc, 'Log transformation: Target variable log-transformed to reduce skewness.')

doc.add_heading('4.3 Feature Engineering', level=2)
doc.add_paragraph('The cleaned dataset is further processed for modeling:')
add_numbered(doc, 'Residential filtering: Non-residential transactions removed (94,435 \u2192 80,791 rows, 14.4% removed).')
add_numbered(doc, 'Target derivation: Log-transformed value exponentiated back, divided by area to get price_per_sqm. Outliers removed using 1st\u201399th percentile, yielding 79,127 rows (range: 502\u201344,367 AED/sqm, mean: 13,003).')
add_numbered(doc, 'Market-history features: Area-level aggregates per year/quarter \u2014 mean price, median price, transaction count, and momentum.')
add_numbered(doc, 'Target encoding: High-cardinality features (area_name, master_project) encoded using mean target encoding with global mean imputation.')
add_numbered(doc, 'One-hot encoding: Low-cardinality categoricals (property_type, property_sub_type, reg_type, nearest_landmark, nearest_metro, nearest_mall, procedure_group) encoded via Spark ML Pipeline.')
add_numbered(doc, 'Feature assembly: All features combined into a 106-dimensional vector using VectorAssembler.')

doc.add_paragraph()
doc.add_paragraph('Table 2: Feature Categories').runs[0].bold = True

add_table(doc,
    ['Category', 'Count', 'Examples'],
    [
        ['Numerical Base', '6', 'procedure_area, bedrooms, has_parking, is_penthouse, sale_year, sale_quarter'],
        ['Market-History', '4', 'area_mean_price, area_median_price, area_transaction_count, area_momentum'],
        ['Target-Encoded', '2', 'area_name_encoded, master_project_encoded'],
        ['One-Hot Encoded', '7 cols (~94 dims)', 'property_type, procedure_group, nearest_metro, etc.'],
        ['Total', '106', ''],
    ]
)

doc.add_heading('4.4 Algorithm Implementation', level=2)

p = doc.add_paragraph()
run = p.add_run('Gradient Boosting (From Scratch):')
run.bold = True
run.font.name = 'Times New Roman'

add_numbered(doc, 'Initialize all predictions to zero.')
add_numbered(doc, 'For each of 10 iterations: compute residuals, train a DecisionTreeRegressor (max depth 5) on residuals, update predictions with learning_rate \u00d7 tree_output (learning rate = 0.1).')
add_numbered(doc, 'Final prediction is the sum of all scaled tree outputs.')

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('Neural Network MLP (PyTorch):')
run.bold = True
run.font.name = 'Times New Roman'

doc.add_paragraph(
    'Input (106 features) \u2192 Linear(106, 128) + ReLU + Dropout(0.3) \u2192 Linear(128, 64) + ReLU '
    '\u2192 Linear(64, 32) + ReLU \u2192 Linear(32, 1) Output'
)

add_bullet(doc, 'on the first hidden layer prevents overfitting.', 'Dropout (0.3) ')
add_bullet(doc, 'on features and target for training stability.', 'StandardScaler ')
add_bullet(doc, '(lr=0.001) adapts learning rates per parameter.', 'Adam optimizer ')
add_bullet(doc, 'directly optimizes for the regression objective.', 'MSELoss ')
add_bullet(doc, 'Batch size 32, 50 epochs with early stopping (patience=10).')

doc.add_heading('4.5 Training and Validation', level=2)

p = doc.add_paragraph()
run = p.add_run('Spark models: ')
run.bold = True
p.add_run('10-fold cross-validation with deterministic hash-based fold assignment. Spark RDD checkpointing manages memory during iterative training.')

p = doc.add_paragraph()
run = p.add_run('Neural Network: ')
run.bold = True
p.add_run('80/20 train/validation split (seed=42). Feature and target scaling with inverse transformation for evaluation. Automatic GPU detection with CPU fallback.')

doc.add_page_break()

# ============================================================
# 5. RESULTS AND EVALUATION
# ============================================================
doc.add_heading('5. Results and Evaluation', level=1)

doc.add_heading('5.1 Spark Model Performance (10-Fold CV)', level=2)

doc.add_paragraph('Table 3: 10-Fold Cross-Validation Results').runs[0].bold = True
add_table(doc,
    ['Model', 'Avg RMSE (AED/sqm)', 'Std RMSE', 'Avg MAE (AED/sqm)', 'Std MAE'],
    [
        ['Baseline (Single Tree)', '3,710.28', '96.93', '2,363.60', '44.21'],
        ['Bagging Ensemble', '3,273.04', '69.23', '2,139.71', '30.78'],
        ['Gradient Boosting', '5,732.31', '66.81', '4,395.89', '45.50'],
    ]
)

doc.add_paragraph('Key observations:')
add_bullet(doc, 'achieved the best RMSE at 3,273.04 AED/sqm \u2014 an 11.78% improvement over the baseline.', 'Bagging ')
add_bullet(doc, 'had the highest RMSE but the lowest standard deviation (66.81), indicating the most consistent predictions across folds. The higher error suggests the conservative learning rate and limited estimators led to underfitting.', 'Gradient Boosting ')
add_bullet(doc, 'decreased with ensemble complexity: baseline Std RMSE 96.93 \u2192 Bagging 69.23 \u2192 Gradient Boosting 66.81.', 'Variance ')

doc.add_heading('5.2 Neural Network MLP Performance', level=2)

doc.add_paragraph('Table 4: Neural Network Evaluation (Train/Val Split)').runs[0].bold = True
add_table(doc,
    ['Metric', 'Value'],
    [
        ['RMSE', '3,810.88 AED/sqm'],
        ['MAE', '2,735.13 AED/sqm'],
        ['R\u00b2 Score', '0.7625'],
    ]
)

doc.add_paragraph(
    'The Neural Network explains 76.25% of the variance in property prices, exceeding the project target '
    'of 0.70. Given the mean price of 13,003 AED/sqm, the MAE represents approximately 21% mean absolute '
    'percentage error.'
)

doc.add_heading('5.3 Final Model Comparison', level=2)

doc.add_paragraph('Table 5: Comprehensive Comparison').runs[0].bold = True
add_table(doc,
    ['Model', 'RMSE', 'MAE', 'R\u00b2', 'Evaluation', 'Framework'],
    [
        ['Baseline Decision Tree', '3,710.28', '2,363.60', '\u2014', '10-Fold CV', 'Spark ML'],
        ['Bagging Ensemble', '3,273.04', '2,139.71', '\u2014', '10-Fold CV', 'Spark ML'],
        ['Gradient Boosting', '5,732.31', '4,395.89', '\u2014', '10-Fold CV', 'Spark ML'],
        ['Neural Network MLP', '3,810.88', '2,735.13', '0.7625', 'Train/Val', 'PyTorch'],
    ]
)

doc.add_heading('5.4 Business Impact', level=2)
add_bullet(doc, 'The Bagging Ensemble offers the best raw accuracy for Spark-deployed scenarios.')
add_bullet(doc, 'The Neural Network provides the strongest explanatory power through its R\u00b2 score.')
add_bullet(doc, 'For a typical 100 sqm apartment valued at ~1,300,000 AED, the MAE translates to approximately 273,500 AED deviation \u2014 useful for ranking and market analysis, though point estimates require further refinement for formal appraisals.')

doc.add_page_break()

# ============================================================
# 6. LIMITATIONS
# ============================================================
doc.add_heading('6. Limitations, Challenges, and Constraints', level=1)

p = doc.add_paragraph()
run = p.add_run('Technical limitations:')
run.bold = True
add_bullet(doc, 'Market-history features are derived from the same dataset, creating a risk of subtle information leakage through area-level aggregates.')
add_bullet(doc, 'The dataset lacks precise coordinates (latitude/longitude), preventing spatial modeling techniques.')
add_bullet(doc, 'The model is trained on a static snapshot and does not account for real-time market shifts.')

p = doc.add_paragraph()
run = p.add_run('Resource constraints:')
run.bold = True
add_bullet(doc, 'Docker environment limited to 4\u20138 GB RAM. A 5% sample (~79,000 rows) was used for local execution.')
add_bullet(doc, 'The from-scratch Gradient Boosting was limited to 10 estimators. Production implementations (XGBoost, LightGBM) with hundreds of trees would perform significantly better.')
add_bullet(doc, 'Extensive Neural Network hyperparameter tuning was not conducted due to time constraints.')

p = doc.add_paragraph()
run = p.add_run('Generalization challenges:')
run.bold = True
add_bullet(doc, 'The model is trained exclusively on Dubai data and may not transfer to other markets.')
add_bullet(doc, 'Structural market changes (regulatory shifts, economic crises) could degrade accuracy over time.')

p = doc.add_paragraph()
run = p.add_run('Ethical considerations:')
run.bold = True
add_bullet(doc, 'If historical data reflects systemic pricing biases, the model may perpetuate them.')
add_bullet(doc, 'Neural networks are less interpretable than tree-based models; stakeholders may require explainability mechanisms (e.g., SHAP values) not currently implemented.')

doc.add_page_break()

# ============================================================
# 7. CONCLUSIONS
# ============================================================
doc.add_heading('7. Conclusions and Discussion', level=1)

doc.add_heading('7.1 Summary of Findings', level=2)
doc.add_paragraph(
    'This project successfully developed an end-to-end AI pipeline for predicting residential property '
    'prices in Dubai. Key findings:'
)
add_numbered(doc, 'Feature engineering is critical: The 106-feature vector incorporating market-history aggregates, target encoding, and one-hot encoding provided a rich representation for effective modeling.')
add_numbered(doc, 'Bagging outperforms other Spark models: The Bagging Ensemble achieved the lowest RMSE (3,273.04 AED/sqm), showing that variance reduction through bootstrap aggregation is most beneficial for this dataset.')
add_numbered(doc, 'Neural Networks achieve strong explanatory power: The MLP achieved R\u00b2 = 0.7625, confirming deep learning can effectively model complex real estate pricing dynamics on tabular data.')
add_numbered(doc, 'Gradient Boosting offers maximum consistency: Despite the highest RMSE, it had the lowest prediction variance (Std RMSE = 66.81).')

doc.add_heading('7.2 Achievement of Objectives', level=2)

doc.add_paragraph('Table 6: Objective Achievement').runs[0].bold = True
add_table(doc,
    ['Objective', 'Status'],
    [
        ['End-to-end data pipeline', 'Achieved'],
        ['Compare ensemble methods with deep learning', 'Achieved \u2014 four models compared'],
        ['R\u00b2 score above 0.70', 'Achieved \u2014 R\u00b2 = 0.7625'],
        ['Minimize RMSE and MAE', 'Achieved \u2014 Bagging RMSE of 3,273'],
        ['Reproducible Docker execution', 'Achieved'],
    ]
)

doc.add_heading('7.3 Key Lessons Learned', level=2)
add_bullet(doc, 'From-scratch implementations provide deep algorithmic understanding but sacrifice the optimizations of mature libraries.')
add_bullet(doc, 'Spark is essential for preprocessing 1.6M rows but introduces overhead for iterative algorithms where in-memory frameworks are more efficient.')
add_bullet(doc, 'Domain knowledge drives feature engineering \u2014 features like area momentum and bedroom extraction required real estate understanding that no automated feature selection could replicate.')

doc.add_page_break()

# ============================================================
# 8. RECOMMENDATIONS
# ============================================================
doc.add_heading('8. Recommendations and Future Work', level=1)

p = doc.add_paragraph()
run = p.add_run('For implementation:')
run.bold = True
add_bullet(doc, 'Deploy the Bagging Ensemble where Spark infrastructure is available for best accuracy.')
add_bullet(doc, 'Deploy the Neural Network where R\u00b2 and explanatory power are prioritized.')
add_bullet(doc, 'Combine both in an ensemble meta-learner (stacking) to leverage complementary strengths.')

p = doc.add_paragraph()
run = p.add_run('Proposed enhancements:')
run.bold = True
add_bullet(doc, 'Replace from-scratch Gradient Boosting with XGBoost/LightGBM using hundreds of trees.')
add_bullet(doc, 'Incorporate latitude/longitude for spatial modeling.')
add_bullet(doc, 'Add time-series components (LSTM) to capture market trends.')
add_bullet(doc, 'Use Bayesian optimization (Optuna) for systematic hyperparameter search.')
add_bullet(doc, 'Integrate SHAP values for per-prediction explainability.')

p = doc.add_paragraph()
run = p.add_run('Maintenance strategy:')
run.bold = True
add_bullet(doc, 'Quarterly retraining with new transaction data.')
add_bullet(doc, 'Monitor prediction error distributions to detect model drift.')
add_bullet(doc, 'Implement schema checks on incoming data.')

p = doc.add_paragraph()
run = p.add_run('Broader applications:')
run.bold = True
add_bullet(doc, 'Adapt the pipeline for other UAE emirates.')
add_bullet(doc, 'Extend to commercial real estate segments.')
add_bullet(doc, 'Wrap the model in a REST API for on-demand valuation queries.')

doc.add_page_break()

# ============================================================
# 9. REFERENCES
# ============================================================
doc.add_heading('9. References', level=1)

refs = [
    'Chen, T., & Guestrin, C. (2016). XGBoost: A scalable tree boosting system. Proceedings of the 22nd ACM SIGKDD, 785\u2013794.',
    'Friedman, J. H. (2001). Greedy function approximation: A gradient boosting machine. Annals of Statistics, 29(5), 1189\u20131232.',
    'Kadra, A., Lindauer, M., Hutter, F., & Grabocka, J. (2021). Well-tuned simple nets excel on tabular datasets. NeurIPS, 34, 23928\u201323941.',
    'Kok, N., Koponen, E. L., & Mart\u00ednez-Barbosa, C. A. (2014). Big data in real estate? The Journal of Portfolio Management, 43(6), 202\u2013211.',
    'Limsombunchai, V. (2004). House price prediction: Hedonic price model vs. artificial neural network. NZARES Conference.',
    'Rosen, S. (1974). Hedonic prices and implicit markets. Journal of Political Economy, 82(1), 34\u201355.',
    'Selim, H. (2009). Determinants of house prices in Turkey. Expert Systems with Applications, 36(2), 2843\u20132852.',
    'Zaharia, M. et al. (2016). Apache Spark: A unified engine for big data processing. Communications of the ACM, 59(11), 56\u201365.',
    'Apache Spark Documentation. (2024). https://spark.apache.org/docs/3.5.0/ml-guide.html',
    'PyTorch Documentation. (2024). https://pytorch.org/tutorials/',
]

for ref in refs:
    p = doc.add_paragraph(ref)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(1.27)
    p.paragraph_format.first_line_indent = Cm(-1.27)

doc.add_page_break()

# ============================================================
# 10. APPENDICES
# ============================================================
doc.add_heading('10. Appendices', level=1)

doc.add_heading('Appendix A: Project File Structure', level=2)

structure = (
    'Project/\n'
    '\u251c\u2500\u2500 Dockerfile\n'
    '\u251c\u2500\u2500 requirements.txt\n'
    '\u251c\u2500\u2500 run_pipeline.sh\n'
    '\u251c\u2500\u2500 PricePulseAI.py              # Preprocessing (PySpark)\n'
    '\u251c\u2500\u2500 notebooks/\n'
    '\u2502   \u2514\u2500\u2500 PricePulseAI.ipynb       # Modeling (Spark ML + PyTorch)\n'
    '\u251c\u2500\u2500 data/\n'
    '\u2502   \u251c\u2500\u2500 Transactions.csv         # Raw (1.6M rows, 279 MB)\n'
    '\u2502   \u2514\u2500\u2500 Transactions_copy.csv/   # Cleaned output\n'
    '\u251c\u2500\u2500 checkpoints/                 # Spark RDD checkpoints\n'
    '\u2514\u2500\u2500 out/                         # Executed outputs'
)
p = doc.add_paragraph(structure)
p.style = doc.styles['Normal']
for run in p.runs:
    run.font.name = 'Courier New'
    run.font.size = Pt(9)

doc.add_heading('Appendix B: Dataset Summary', level=2)
add_table(doc,
    ['Attribute', 'Value'],
    [
        ['Total raw records', '1,644,775'],
        ['Raw columns', '46'],
        ['After residential filtering', '80,791'],
        ['After outlier removal', '79,127'],
        ['Price range', '502\u201344,367 AED/sqm'],
        ['Mean price', '13,003 AED/sqm'],
        ['Median price', '11,405 AED/sqm'],
        ['Final feature count', '106'],
    ]
)

# ============================================================
# SAVE
# ============================================================
output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'CSCI323_Project_Report.docx')
doc.save(output_path)
print(f"Report saved to: {output_path}")
